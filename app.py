# **************************************************
import os
import docx

os.system(command="cls")

print("Version of 'docx' library:", docx.__version__)
# **************************************************


# **************************************************
# import os
# from docx import Document

# os.system(command="cls")

# file_path: str = "./files/test_01.docx"
# text: str = "This is the first paragraph."

# document = Document()
# paragraph = document.add_paragraph(text=text)
# document.save(path_or_stream=file_path)
# **************************************************


# **************************************************
# import os
# from docx import Document

# os.system(command="cls")

# file_path: str = "./files/test_01.docx"

# document = Document(docx=file_path)
# paragraph_count: int = len(document.paragraphs)
# print(f"Paragraph Count: {paragraph_count}")

# for paragraph in document.paragraphs:
#     print(paragraph.text)
#     print()
# **************************************************


# **************************************************
# import os
# from docx import Document
# from docx.shared import Pt
# from docx.shared import RGBColor
# from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.text import WD_ALIGN_PARAGRAPH

# os.system(command="cls")

# file_path: str = "./files/test_02.docx"

# document = Document()

# LTR_HEADER_STYLE = document.styles.add_style("ltr_header", WD_STYLE_TYPE.PARAGRAPH)
# LTR_HEADER_STYLE.base_style = document.styles["Normal"]
# LTR_HEADER_STYLE.font.bold = True
# LTR_HEADER_STYLE.font.name = "Verdana"
# LTR_HEADER_STYLE.font.size = Pt(points=16)
# LTR_HEADER_STYLE.font.color.rgb = RGBColor(255, 0, 0)

# LTR_NORMAL_STYLE = document.styles.add_style("ltr_normal", WD_STYLE_TYPE.PARAGRAPH)
# LTR_NORMAL_STYLE.base_style = document.styles["Normal"]
# LTR_NORMAL_STYLE.font.bold = False
# LTR_NORMAL_STYLE.font.name = "Verdana"
# LTR_NORMAL_STYLE.font.size = Pt(points=12)
# LTR_NORMAL_STYLE.font.color.rgb = RGBColor(0, 0, 255)

# text: str

# text = """
# In the name of God
# """.strip()
# paragraph = document.add_paragraph(text=text, style=LTR_HEADER_STYLE)
# paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# text = """
# This is the first paragraph.
# """.strip()
# paragraph = document.add_paragraph(text=text, style=LTR_NORMAL_STYLE)
# paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

# document.save(path_or_stream=file_path)
# **************************************************


# **************************************************
# import os
# from docx import Document
# from docx.oxml.ns import qn

# # from docx.shared import Pt
# from docx.text.run import Run
# from docx.shared import RGBColor
# from docx.oxml import OxmlElement
# from docx.text.paragraph import Paragraph
# from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.text import WD_ALIGN_PARAGRAPH


# def fix_paragraph_cs_formatting(
#     paragraph: Paragraph,
# ):
#     """
#     Fix Paragraph CS Formatting
#     CS: Complex Script, ex, arabic
#     """

#     rpr = paragraph._element.get_or_add_pPr()

#     lang = OxmlElement("w:lang")  # Set Language
#     lang.set(qn("w:bidi"), "fa-IR")
#     rpr.append(lang)

#     bidi = OxmlElement("w:bidi")
#     rpr.append(bidi)


# def fix_run_cs_formatting(
#     run: Run,
#     font_size: int,
#     font_name: str,
#     bold: bool = False,
# ):
#     """
#     Fix Run CS Formatting
#     CS: Complex Script, ex, arabic
#     """

#     font_size_str: str = str(font_size * 2)

#     rpr = run.element.get_or_add_rPr()

#     sz = OxmlElement("w:sz")  # Set font size
#     sz.set(qn("w:val"), font_size_str)
#     rpr.append(sz)

#     szcs = OxmlElement("w:szCs")  # Set font size
#     szcs.set(qn("w:val"), font_size_str)
#     rpr.append(szcs)

#     rFonts = rpr.get_or_add_rFonts()

#     rFonts.set(qn("w:cs"), font_name)  # Set font name
#     rFonts.set(qn("w:ascii"), font_name)  # Set font name
#     rFonts.set(qn("w:hAnsi"), font_name)  # Set font name

#     if bold:
#         b = OxmlElement("w:b")  # Set bold for the english language
#         b.set(qn("w:val"), "True")
#         rpr.append(b)

#         bcs = OxmlElement("w:bCs")  # Set bold for the complex language
#         bcs.set(qn("w:val"), "True")
#         rpr.append(bcs)


# def add_persian_header(document, text: str) -> None:
#     """
#     Add Persian Header
#     """

#     paragraph = document.add_paragraph(style=RTL_HEADER_STYLE)
#     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     fix_paragraph_cs_formatting(paragraph=paragraph)
#     run = paragraph.add_run(text=text)
#     fix_run_cs_formatting(
#         run=run,
#         bold=True,
#         font_size=16,
#         font_name="IRANSansX",
#     )


# def add_persian_paragraph(document, text: str) -> None:
#     """
#     Add Persian Paragraph
#     """

#     paragraph = document.add_paragraph(style=RTL_NORMAL_STYLE)
#     paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
#     fix_paragraph_cs_formatting(paragraph=paragraph)
#     run = paragraph.add_run(text=text)
#     fix_run_cs_formatting(
#         run=run,
#         bold=False,
#         font_size=12,
#         font_name="IRANSansX",
#     )


# os.system(command="cls")

# file_path: str = "./files/test_03.docx"

# document = Document()

# RTL_HEADER_STYLE = document.styles.add_style("rtl_header", WD_STYLE_TYPE.PARAGRAPH)
# RTL_HEADER_STYLE.font.rtl = True
# RTL_HEADER_STYLE.font.complex_script = True
# RTL_HEADER_STYLE.font.color.rgb = RGBColor(255, 0, 0)
# RTL_HEADER_STYLE.base_style = document.styles["Normal"]

# RTL_NORMAL_STYLE = document.styles.add_style("rtl_normal", WD_STYLE_TYPE.PARAGRAPH)
# RTL_NORMAL_STYLE.font.rtl = True
# RTL_NORMAL_STYLE.font.complex_script = True
# RTL_NORMAL_STYLE.font.color.rgb = RGBColor(0, 0, 255)
# RTL_NORMAL_STYLE.base_style = document.styles["Normal"]

# text: str

# text = """
# به نام آن‌که جان را فکرت آموخت
# """.strip()
# paragraph = document.add_paragraph(style=RTL_HEADER_STYLE)
# paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
# fix_paragraph_cs_formatting(paragraph=paragraph)
# run = paragraph.add_run(text=text)
# fix_run_cs_formatting(
#     run=run,
#     bold=True,
#     font_size=16,
#     font_name="IRANSansX",
# )

# text = """
# به گزارش خبرنگار سیاسی خبرگزاری فارس، سیدمحمد حسینی عضو هیئت علمی دانشگاه و معاون پارلمانی دولت شهید رئیسی در پیامی به شدت به رویکرد دوگانه غرب و سازمان‌های بین‌المللی وابسته به آن‌ها واکنش نشان داد. وی با اشاره به جنایات اخیر رژیم صهیونیستی در غزه و رویکرد سازمان‌های بین‌المللی، تاکید کرد: «در حالی که اسرائیل خون‌آشام با نقض آتش‌بس در یک سحرگاه بیش از ۱۰۰۰ نفر را در غزه شهید و مجروح می‌کند، شورای حقوق بشر سازمان ملل ایران را به دلیل بازداشت یا حکم اعدام چند جنایتکار ناقض حقوق بشر معرفی می‌کند! راستی مفهوم واژه‌ها چقدر لوث و واژگونه شده است!»
# """.strip()
# paragraph = document.add_paragraph(style=RTL_NORMAL_STYLE)
# paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
# fix_paragraph_cs_formatting(paragraph=paragraph)
# run = paragraph.add_run(text=text)
# fix_run_cs_formatting(
#     run=run,
#     bold=False,
#     font_size=12,
#     font_name="IRANSansX",
# )

# text = """
# به نام آن‌که جان را فکرت آموخت
# """.strip()
# add_persian_header(document=document, text=text)

# text = """
# به گزارش خبرنگار سیاسی خبرگزاری فارس، سیدمحمد حسینی عضو هیئت علمی دانشگاه و معاون پارلمانی دولت شهید رئیسی در پیامی به شدت به رویکرد دوگانه غرب و سازمان‌های بین‌المللی وابسته به آن‌ها واکنش نشان داد. وی با اشاره به جنایات اخیر رژیم صهیونیستی در غزه و رویکرد سازمان‌های بین‌المللی، تاکید کرد: «در حالی که اسرائیل خون‌آشام با نقض آتش‌بس در یک سحرگاه بیش از ۱۰۰۰ نفر را در غزه شهید و مجروح می‌کند، شورای حقوق بشر سازمان ملل ایران را به دلیل بازداشت یا حکم اعدام چند جنایتکار ناقض حقوق بشر معرفی می‌کند! راستی مفهوم واژه‌ها چقدر لوث و واژگونه شده است!»
# """.strip()
# add_persian_paragraph(document=document, text=text)

# document.save(path_or_stream=file_path)
# **************************************************


# **************************************************
# import os
# from docx import Document
# from docx.shared import Pt
# from docx.oxml.ns import qn
# from docx.text.run import Run
# from docx.shared import RGBColor
# from docx.oxml import OxmlElement
# from docx.text.paragraph import Paragraph
# from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.text import WD_ALIGN_PARAGRAPH


# def fix_paragraph_cs_formatting(
#     paragraph: Paragraph,
# ):
#     """
#     Fix Paragraph CS Formatting
#     CS: Complex Script, ex, arabic
#     """

#     rpr = paragraph._element.get_or_add_pPr()

#     lang = OxmlElement("w:lang")  # Set Language
#     lang.set(qn("w:bidi"), "fa-IR")
#     rpr.append(lang)

#     bidi = OxmlElement("w:bidi")
#     rpr.append(bidi)


# def fix_run_cs_formatting(
#     run: Run,
#     font_size: int,
#     font_name: str,
#     bold: bool = False,
# ):
#     """
#     Fix Run CS Formatting
#     CS: Complex Script, ex, arabic
#     """

#     font_size_str: str = str(font_size * 2)

#     rpr = run.element.get_or_add_rPr()

#     sz = OxmlElement("w:sz")  # Set font size
#     sz.set(qn("w:val"), font_size_str)
#     rpr.append(sz)

#     szcs = OxmlElement("w:szCs")  # Set font size
#     szcs.set(qn("w:val"), font_size_str)
#     rpr.append(szcs)

#     rFonts = rpr.get_or_add_rFonts()

#     rFonts.set(qn("w:cs"), font_name)  # Set font name
#     rFonts.set(qn("w:ascii"), font_name)  # Set font name
#     rFonts.set(qn("w:hAnsi"), font_name)  # Set font name

#     if bold:
#         b = OxmlElement("w:b")  # Set bold for the english language
#         b.set(qn("w:val"), "True")
#         rpr.append(b)

#         bcs = OxmlElement("w:bCs")  # Set bold for the complex language
#         bcs.set(qn("w:val"), "True")
#         rpr.append(bcs)


# def add_persian_header(document, text: str) -> None:
#     """
#     Add Persian Header
#     """

#     paragraph = document.add_paragraph(style=RTL_HEADER_STYLE)
#     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     fix_paragraph_cs_formatting(paragraph=paragraph)
#     run = paragraph.add_run(text=text)
#     fix_run_cs_formatting(
#         run=run,
#         bold=True,
#         font_size=16,
#         font_name="IRANSansX",
#     )


# def add_persian_paragraph(document, text: str) -> None:
#     """
#     Add Persian Paragraph
#     """

#     paragraph = document.add_paragraph(style=RTL_NORMAL_STYLE)
#     paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
#     fix_paragraph_cs_formatting(paragraph=paragraph)
#     run = paragraph.add_run(text=text)
#     fix_run_cs_formatting(
#         run=run,
#         bold=False,
#         font_size=12,
#         font_name="IRANSansX",
#     )


# def add_english_header(document, text: str) -> None:
#     """
#     Add English Header
#     """

#     paragraph = document.add_paragraph(text=text, style=LTR_HEADER_STYLE)
#     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


# def add_english_paragraph(document, text: str) -> None:
#     """
#     Add English Paragraph
#     """

#     paragraph = document.add_paragraph(text=text, style=LTR_NORMAL_STYLE)
#     paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


# os.system(command="cls")

# file_path: str = "./files/test_04.docx"

# document = Document()

# LTR_HEADER_STYLE = document.styles.add_style("ltr_header", WD_STYLE_TYPE.PARAGRAPH)
# LTR_HEADER_STYLE.base_style = document.styles["Normal"]
# LTR_HEADER_STYLE.font.bold = True
# LTR_HEADER_STYLE.font.name = "Verdana"
# LTR_HEADER_STYLE.font.size = Pt(points=16)
# LTR_HEADER_STYLE.font.color.rgb = RGBColor(255, 0, 0)

# LTR_NORMAL_STYLE = document.styles.add_style("ltr_normal", WD_STYLE_TYPE.PARAGRAPH)
# LTR_NORMAL_STYLE.base_style = document.styles["Normal"]
# LTR_NORMAL_STYLE.font.bold = False
# LTR_NORMAL_STYLE.font.name = "Verdana"
# LTR_NORMAL_STYLE.font.size = Pt(points=12)
# LTR_NORMAL_STYLE.font.color.rgb = RGBColor(0, 0, 255)

# RTL_HEADER_STYLE = document.styles.add_style("my_header", WD_STYLE_TYPE.PARAGRAPH)
# RTL_HEADER_STYLE.font.rtl = True
# RTL_HEADER_STYLE.font.complex_script = True
# RTL_HEADER_STYLE.font.color.rgb = RGBColor(255, 0, 0)
# RTL_HEADER_STYLE.base_style = document.styles["Normal"]

# RTL_NORMAL_STYLE = document.styles.add_style("my_normal", WD_STYLE_TYPE.PARAGRAPH)
# RTL_NORMAL_STYLE.font.rtl = True
# RTL_NORMAL_STYLE.font.complex_script = True
# RTL_NORMAL_STYLE.font.color.rgb = RGBColor(0, 0, 255)
# RTL_NORMAL_STYLE.base_style = document.styles["Normal"]

# text: str

# text = """
# به نام آن‌که جان را فکرت آموخت
# """.strip()
# add_persian_header(document=document, text=text)

# text = """
# به گزارش خبرنگار سیاسی خبرگزاری فارس، سیدمحمد حسینی عضو هیئت علمی دانشگاه و معاون پارلمانی دولت شهید رئیسی در پیامی به شدت به رویکرد دوگانه غرب و سازمان‌های بین‌المللی وابسته به آن‌ها واکنش نشان داد. وی با اشاره به جنایات اخیر رژیم صهیونیستی در غزه و رویکرد سازمان‌های بین‌المللی، تاکید کرد: «در حالی که اسرائیل خون‌آشام با نقض آتش‌بس در یک سحرگاه بیش از ۱۰۰۰ نفر را در غزه شهید و مجروح می‌کند، شورای حقوق بشر سازمان ملل ایران را به دلیل بازداشت یا حکم اعدام چند جنایتکار ناقض حقوق بشر معرفی می‌کند! راستی مفهوم واژه‌ها چقدر لوث و واژگونه شده است!»
# """.strip()
# add_persian_paragraph(document=document, text=text)

# text = """
# In the name of God
# """.strip()
# add_english_header(document=document, text=text)

# text = """
# This is the first paragraph.
# """.strip()
# add_english_paragraph(document=document, text=text)

# document.save(path_or_stream=file_path)
# **************************************************


# **************************************************
# import os
# import dt_docx as docx

# os.system(command="cls")

# text: str

# text = """
# به نام آن‌که جان را فکرت آموخت
# """.strip()
# docx.add_persian_header(text=text)

# text = """
# به گزارش خبرنگار سیاسی خبرگزاری فارس، سیدمحمد حسینی عضو هیئت علمی دانشگاه و معاون پارلمانی دولت شهید رئیسی در پیامی به شدت به رویکرد دوگانه غرب و سازمان‌های بین‌المللی وابسته به آن‌ها واکنش نشان داد. وی با اشاره به جنایات اخیر رژیم صهیونیستی در غزه و رویکرد سازمان‌های بین‌المللی، تاکید کرد: «در حالی که اسرائیل خون‌آشام با نقض آتش‌بس در یک سحرگاه بیش از ۱۰۰۰ نفر را در غزه شهید و مجروح می‌کند، شورای حقوق بشر سازمان ملل ایران را به دلیل بازداشت یا حکم اعدام چند جنایتکار ناقض حقوق بشر معرفی می‌کند! راستی مفهوم واژه‌ها چقدر لوث و واژگونه شده است!»
# """.strip()
# docx.add_persian_paragraph(text=text)

# text = """
# In the name of God
# """.strip()
# docx.add_english_header(text=text)

# text = """
# This is the first paragraph.
# """.strip()
# docx.add_english_paragraph(text=text)

# file_path: str = "./files/test_05.docx"
# docx.save(path=file_path)
# **************************************************

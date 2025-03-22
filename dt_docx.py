from docx import Document

from docx.shared import Pt
from docx.shared import RGBColor

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.text.run import Run
from docx.text.paragraph import Paragraph

from docx.styles.style import BaseStyle
from docx.document import Document as DocumentClass


def fix_paragraph_cs_formatting(
    paragraph: Paragraph,
):
    """
    Fix Paragraph CS Formatting Function
    CS: Complex Script, ex, arabic
    """

    rpr = paragraph._element.get_or_add_pPr()

    lang = OxmlElement("w:lang")  # Set Language
    lang.set(qn("w:bidi"), "fa-IR")
    rpr.append(lang)

    bidi = OxmlElement("w:bidi")
    rpr.append(bidi)


def fix_run_cs_formatting(
    run: Run,
    font_size: int,
    font_name: str,
    bold: bool = False,
):
    """
    Fix Run CS Formatting Function
    CS: Complex Script, ex, arabic
    """

    font_size_str: str = str(font_size * 2)

    rpr = run.element.get_or_add_rPr()

    sz = OxmlElement("w:sz")  # Set font size
    sz.set(qn("w:val"), font_size_str)
    rpr.append(sz)

    szcs = OxmlElement("w:szCs")  # Set font size
    szcs.set(qn("w:val"), font_size_str)
    rpr.append(szcs)

    rFonts = rpr.get_or_add_rFonts()

    rFonts.set(qn("w:cs"), font_name)  # Set font name
    rFonts.set(qn("w:ascii"), font_name)  # Set font name
    rFonts.set(qn("w:hAnsi"), font_name)  # Set font name

    if bold:
        b = OxmlElement("w:b")  # Set bold for the english language
        b.set(qn("w:val"), "True")
        rpr.append(b)

        bcs = OxmlElement("w:bCs")  # Set bold for the complex language
        bcs.set(qn("w:val"), "True")
        rpr.append(bcs)


def add_persian_header(text: str) -> None:
    """
    Add Persian Header Function
    """

    paragraph = document.add_paragraph(style=RTL_HEADER_STYLE)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fix_paragraph_cs_formatting(paragraph=paragraph)
    run = paragraph.add_run(text=text)
    fix_run_cs_formatting(
        run=run,
        bold=True,
        font_size=16,
        font_name="IRANSansX",
    )


def add_persian_paragraph(text: str) -> None:
    """
    Add Persian Paragraph Function
    """

    paragraph = document.add_paragraph(style=RTL_NORMAL_STYLE)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fix_paragraph_cs_formatting(paragraph=paragraph)
    run = paragraph.add_run(text=text)
    fix_run_cs_formatting(
        run=run,
        bold=False,
        font_size=12,
        font_name="IRANSansX",
    )


def add_english_header(text: str) -> None:
    """
    Add English Header Function
    """

    paragraph = document.add_paragraph(text=text, style=LTR_HEADER_STYLE)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_english_paragraph(text: str) -> None:
    """
    Add English Paragraph Function
    """

    paragraph = document.add_paragraph(text=text, style=LTR_NORMAL_STYLE)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def save(path: str) -> None:
    """
    Save Function
    """

    document.save(path_or_stream=path)


document: DocumentClass = Document()

LTR_HEADER_STYLE: BaseStyle = document.styles.add_style(
    "ltr_header", WD_STYLE_TYPE.PARAGRAPH
)
LTR_HEADER_STYLE.base_style = document.styles["Normal"]
LTR_HEADER_STYLE.font.bold = True
LTR_HEADER_STYLE.font.name = "Verdana"
LTR_HEADER_STYLE.font.size = Pt(points=16)
LTR_HEADER_STYLE.font.color.rgb = RGBColor(255, 0, 0)

LTR_NORMAL_STYLE: BaseStyle = document.styles.add_style(
    "ltr_normal", WD_STYLE_TYPE.PARAGRAPH
)
LTR_NORMAL_STYLE.base_style = document.styles["Normal"]
LTR_NORMAL_STYLE.font.bold = False
LTR_NORMAL_STYLE.font.name = "Verdana"
LTR_NORMAL_STYLE.font.size = Pt(points=12)
LTR_NORMAL_STYLE.font.color.rgb = RGBColor(0, 0, 255)

RTL_HEADER_STYLE: BaseStyle = document.styles.add_style(
    "my_header", WD_STYLE_TYPE.PARAGRAPH
)
RTL_HEADER_STYLE.font.rtl = True
RTL_HEADER_STYLE.font.complex_script = True
RTL_HEADER_STYLE.font.color.rgb = RGBColor(255, 0, 0)
RTL_HEADER_STYLE.base_style = document.styles["Normal"]

RTL_NORMAL_STYLE: BaseStyle = document.styles.add_style(
    "my_normal", WD_STYLE_TYPE.PARAGRAPH
)
RTL_NORMAL_STYLE.font.rtl = True
RTL_NORMAL_STYLE.font.complex_script = True
RTL_NORMAL_STYLE.font.color.rgb = RGBColor(0, 0, 255)
RTL_NORMAL_STYLE.base_style = document.styles["Normal"]
